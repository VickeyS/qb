import streamlit as st
import requests
import datetime
from PyPDF2 import PdfReader
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import mm
from reportlab.lib.utils import simpleSplit
import io
import os
import re
import sqlite3

# Helper functions
def extract_text_from_pdf(file):
    reader = PdfReader(file)
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(file):
    doc = Document(file)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

def generate_questions_gemini(api_key, material, params):
    prompt = f"""
    Generate a question paper from the following material.\nMaterial:\n{material}\n\nParameters:\nTotal Marks: {params['total_marks']}\nMCQs (1 mark): {params['mcq_count']}\nOne-liner (1 mark): {params['one_liner_count']}\nShort (2 marks): {params['short_count']}\nLong (5 marks): {params['long_count']}\nFormat: List questions with marks.\n"""
    url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key=" + api_key
    headers = {"Content-Type": "application/json"}
    data = {
        "contents": [{"parts": [{"text": prompt}]}]
    }
    response = requests.post(url, headers=headers, json=data)
    # Show raw response for debugging
    print('Gemini API raw response:', response.text)
    if response.status_code == 200:
        try:
            return response.json()['candidates'][0]['content']['parts'][0]['text']
        except Exception:
            return f"Error: Unexpected response format. Raw: {response.text}"
    else:
        return f"Error: {response.text}"

def create_pdf(questions, total_marks, subject="Fundamentals of AI", exam_title="LJ Polytechnic", time="2 Hours"):
    from io import BytesIO
    from reportlab.lib import colors
    from reportlab.lib.utils import simpleSplit
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    # Attractive Header: colored bar, bold fonts
    c.setFillColor(colors.HexColor('#003366'))
    c.rect(0, height-90, width, 70, fill=1, stroke=0)
    c.setFillColor(colors.white)
    c.setFont("Helvetica-Bold", 20)
    c.drawString(35, height-40, exam_title)
    c.setFont("Helvetica", 13)
    c.drawRightString(width-35, height-40, f"Time: {time}")
    # Centered subject in blue bar
    sanitized_subject = subject.split('#')[0].strip()
    c.setFont("Helvetica-Bold", 16)
    subject_width = c.stringWidth(sanitized_subject, "Helvetica-Bold", 16)
    c.drawString((width - subject_width) / 2, height-60, sanitized_subject)
    # Centered total marks in blue bar
    c.setFont("Helvetica-Bold", 13)
    total_marks_text = f"Total Marks: {total_marks}"
    marks_width = c.stringWidth(total_marks_text, "Helvetica-Bold", 13)
    c.drawString((width - marks_width) / 2, height-78, total_marks_text)
    c.setFillColor(colors.black)
    y = height - 105
    # Instructions box
    c.setStrokeColor(colors.HexColor('#003366'))
    c.setLineWidth(1)
    c.roundRect(25, y-40, width-50, 40, 8, stroke=1, fill=0)
    c.setFont("Helvetica-Bold", 12)
    c.drawString(35, y-18, "Instructions:")
    c.setFont("Helvetica", 12)
    c.drawString(55, y-32, "1. Answer all questions.  2. Marks are indicated against each question.")
    y -= 60
    # Section headers and instructions (only once per section)
    section_headers = {
        'A': ("Section - A", "Multiple Choice Questions (10 Marks)", "Attempt all questions. Choose the correct option."),
        'B': ("Section - B", "One-liner Questions (10 Marks) (1 mark each)", None),
        'C': ("Section - C", "Short Questions (10 Marks)", "Answer the following (2 marks each):"),
        'D': ("Section - D", "Descriptive Questions (20 Marks)", "Answer any 4 out of the following 5 questions (5 marks each):")
    }
    section = 'A'
    q_num = 1
    in_mcq = True
    in_very_short = False
    in_short = False
    in_desc = False
    shown_sections = set()
    questions = questions.replace('\r\n', '\n').replace('\r', '\n')
    lines = questions.split('\n')
    mcq_pattern = re.compile(r'^(\d+)[\).:\- ]+(.+)')
    # --- SEGMENT QUESTIONS INTO SECTIONS ---
    # We'll use simple heuristics to segment questions by marks and type
    section_a = []  # MCQs (1 mark)
    section_b = []  # One-liners (1 mark)
    section_c = []  # Short (2 marks)
    section_d = []  # Long/Descriptive (5 marks)
    current_section = None
    for line in lines:
        l = line.strip()
        if not l:
            continue
        # Heuristic: MCQ options (a), b), etc.)
        if l and (l[0] in 'abcd' and (l[1:3] == ') ' or l[1:2] == '.' or l[1:2] == '-')):
            if current_section == 'A' and section_a:
                section_a[-1].append(l)
            continue
        # MCQ question (1 mark)
        if re.match(r"^\d+[\).:\- ]+.*", l) and 'mark' not in l.lower() and 'short' not in l.lower() and 'long' not in l.lower():
            section_a.append([l])
            current_section = 'A'
            continue
        # One-liner (1 mark)
        if ('one-liner' in l.lower() or 'very short' in l.lower() or '1 mark' in l.lower()) and not l.lower().startswith('section'):
            section_b.append(l)
            current_section = 'B'
            continue
        # Short (2 marks)
        if ('2 mark' in l.lower() or 'short' in l.lower()) and not l.lower().startswith('section'):
            section_c.append(l)
            current_section = 'C'
            continue
        # Long/Descriptive (5 marks)
        if ('5 mark' in l.lower() or 'long' in l.lower() or 'descriptive' in l.lower()) and not l.lower().startswith('section'):
            section_d.append(l)
            current_section = 'D'
            continue
        # Fallback: if in a section, treat as question
        if current_section == 'A':
            section_a[-1].append(l)
        elif current_section == 'B':
            section_b.append(l)
        elif current_section == 'C':
            section_c.append(l)
        elif current_section == 'D':
            section_d.append(l)
    # --- RENDER SECTIONS IN ORDER WITH SPACING AND CLEAR SEGMENTATION ---
    def render_section(title, subtitle, instructions, questions, is_mcq=False, y=y):
        nonlocal c
        # Add extra space before each section for clarity
        y -= 25
        if y < 100:
            c.showPage()
            y = height - 60
        c.setFillColor(colors.HexColor('#003366'))
        c.setFont("Helvetica-Bold", 15)
        c.drawString(30, y, title)
        y -= 20
        c.setFillColor(colors.black)
        c.setFont("Helvetica-Bold", 12)
        c.drawString(30, y, subtitle)
        y -= 16
        if instructions:
            c.setFont("Helvetica-Oblique", 11)
            c.drawString(30, y, instructions)
            y -= 16
        # Render questions
        for q in questions:
            if is_mcq:
                # q is a list: [question, option1, option2, ...]
                wrapped = simpleSplit(q[0], "Helvetica", 12, 500)
                for wline in wrapped:
                    if y < 60:
                        c.showPage()
                        c.setFont("Helvetica", 12)
                        y = height - 40
                    c.setFont("Helvetica", 12)
                    c.drawString(30, y, wline)
                    y -= 16
                for opt in q[1:]:
                    wrapped_opt = simpleSplit(opt, "Helvetica", 12, 470)
                    for wline in wrapped_opt:
                        if y < 40:
                            c.showPage()
                            c.setFont("Helvetica", 12)
                            y = height - 40
                        c.drawString(70, y, wline)
                        y -= 14
            else:
                wrapped = simpleSplit(q, "Helvetica", 12, 500)
                for wline in wrapped:
                    if y < 60:
                        c.showPage()
                        c.setFont("Helvetica", 12)
                        y = height - 40
                    c.setFont("Helvetica", 12)
                    c.drawString(30, y, wline)
                    y -= 18
        # Add extra space after each section
        y -= 10
        return y
    # Render all sections in order, with clear spacing
    y = render_section("Section - A", "Multiple Choice Questions (10 Marks)", "Attempt all questions. Choose the correct option.", section_a, is_mcq=True, y=y)
    y = render_section("Section - B", "One-liner Questions (10 Marks) (1 mark each)", None, section_b, y=y)
    y = render_section("Section - C", "Short Questions (10 Marks)", "Answer the following (2 marks each):", section_c, y=y)
    y = render_section("Section - D", "Descriptive Questions (20 Marks)", "Answer any 4 out of the following 5 questions (5 marks each):", section_d, y=y)
    c.save()
    buffer.seek(0)
    return buffer

def create_pdf_from_inputs(mcqs, one_liners, shorts, longs, total_marks, subject="Fundamentals of AI", exam_title="LJ Polytechnic", time="2 Hours"):
    from io import BytesIO
    from reportlab.lib import colors
    from reportlab.lib.utils import simpleSplit
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    # Attractive Header: colored bar, bold fonts
    c.setFillColor(colors.HexColor('#003366'))
    c.rect(0, height-70, width, 50, fill=1, stroke=0)
    c.setFillColor(colors.white)
    c.setFont("Helvetica-Bold", 20)
    c.drawString(35, height-40, exam_title)
    c.setFont("Helvetica", 13)
    c.drawRightString(width-35, height-40, f"Time: {time}")
    # Centered subject in blue bar
    sanitized_title = subject.split('#')[0].strip()
    c.setFont("Helvetica-Bold", 16)
    title_width = c.stringWidth(sanitized_title, "Helvetica-Bold", 16)
    c.drawString((width - title_width) / 2, height-60, sanitized_title)
    c.setFillColor(colors.black)
    y = height - 95
    # Instructions box
    c.setStrokeColor(colors.HexColor('#003366'))
    c.setLineWidth(1)
    c.roundRect(25, y-40, width-50, 40, 8, stroke=1, fill=0)
    c.setFont("Helvetica-Bold", 12)
    c.drawString(35, y-18, "Instructions:")
    c.setFont("Helvetica", 12)
    c.drawString(55, y-32, "1. Answer all questions.  2. Marks are indicated against each question.")
    y -= 70
    def render_section(title, subtitle, instructions, questions, is_mcq=False, y=y):
        nonlocal c
        # Add extra vertical space before section
        y -= 20
        if y < 120:
            c.showPage()
            y = height - 60
        c.setFillColor(colors.HexColor('#003366'))
        c.setFont("Helvetica-Bold", 14)
        c.drawString(30, y, title)
        y -= 18
        c.setFillColor(colors.black)
        c.setFont("Helvetica", 12)
        c.drawString(30, y, subtitle)
        y -= 16
        if instructions:
            c.setFont("Helvetica-Oblique", 11)
            c.drawString(30, y, instructions)
            y -= 16
        # Render questions
        q_num = 1
        for q in questions:
            if is_mcq:
                # q: question + options (split by lines)
                lines = q.strip().split('\n')
                if not lines:
                    continue
                # Question
                wrapped = simpleSplit(f"{q_num}. {lines[0]}", "Helvetica", 12, 500)
                for wline in wrapped:
                    if y < 60:
                        c.showPage()
                        c.setFont("Helvetica", 12)
                        y = height - 40
                    c.drawString(30, y, wline)
                    y -= 16
                # Options
                for opt in lines[1:]:
                    wrapped_opt = simpleSplit(opt, "Helvetica", 12, 470)
                    for wline in wrapped_opt:
                        if y < 40:
                            c.showPage()
                            c.setFont("Helvetica", 12)
                            y = height - 40
                        c.drawString(70, y, wline)
                        y -= 14
                q_num += 1
            else:
                wrapped = simpleSplit(f"{q_num}. {q.strip()}", "Helvetica", 12, 500)
                for wline in wrapped:
                    if y < 60:
                        c.showPage()
                        c.setFont("Helvetica", 12)
                        y = height - 40
                    c.drawString(30, y, wline)
                    y -= 18
                q_num += 1
        # Add extra vertical space after section
        return y - 20
    # Render all sections in order
    y = render_section("Section - A", "Multiple Choice Questions (10 Marks)", "Attempt all questions. Choose the correct option.", mcqs, is_mcq=True, y=y)
    y = render_section("Section - B", "One-liner Questions (10 Marks) (1 mark each)", None, one_liners, y=y)
    y = render_section("Section - C", "Short Questions (10 Marks)", "Answer the following (2 marks each):", shorts, y=y)
    y = render_section("Section - D", "Descriptive Questions (20 Marks)", "Answer any 4 out of the following 5 questions (5 marks each):", longs, y=y)
    c.save()
    buffer.seek(0)
    return buffer

def create_pdf_from_inputs_preserve_format(raw_text, file_title="Question Paper", exam_title="LJ Polytechnic", time="2 Hours", total_marks=None):
    from io import BytesIO
    from reportlab.lib import colors
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    # Header (blue bar with centered subject and total marks)
    c.setFillColor(colors.HexColor('#003366'))
    c.rect(0, height-90, width, 70, fill=1, stroke=0)
    c.setFillColor(colors.white)
    c.setFont("Helvetica-Bold", 20)
    c.drawString(35, height-40, exam_title)
    c.setFont("Helvetica", 13)
    c.drawRightString(width-35, height-40, f"Time: {time}")
    # Centered subject in blue bar
    sanitized_title = file_title.split('#')[0].strip()
    c.setFont("Helvetica-Bold", 16)
    title_width = c.stringWidth(sanitized_title, "Helvetica-Bold", 16)
    c.drawString((width - title_width) / 2, height-60, sanitized_title)
    # Centered total marks in blue bar (if provided)
    if total_marks is not None:
        c.setFont("Helvetica-Bold", 13)
        total_marks_text = f"Total Marks: {total_marks}"
        marks_width = c.stringWidth(total_marks_text, "Helvetica-Bold", 13)
        c.drawString((width - marks_width) / 2, height-78, total_marks_text)
    c.setFillColor(colors.black)
    y = height - 105
    # Remove unwanted lines from raw_text
    lines = raw_text.replace('**', '').replace('##', '').split('\n')
    skip_patterns = [
        r"^lj polytechnic$",
        r"^fundamentals of ai.*question paper$"
    ]
    import re
    filtered_lines = []
    for line in lines:
        l = line.strip().lower()
        if any(re.match(pat, l) for pat in skip_patterns):
            continue
        filtered_lines.append(line)
    # Use Times New Roman for all body lines, wrap text to fit page
    font_name = "Times-Roman"
    c.setFont(font_name, 12)
    max_width = width - 60  # 30pt margin left/right
    section_header_pattern = re.compile(r"^\s*Section\s*[-:]?\s*[A-Z][\s:.-]*", re.IGNORECASE)
    for line in filtered_lines:
        # Detect section header (e.g., Section A:, Section - A, Section A)
        if section_header_pattern.match(line.strip()):
            # Center and bold section header
            if y < 60:
                c.showPage()
                c.setFont("Helvetica-Bold", 14)
                y = height - 40
            c.setFont("Helvetica-Bold", 14)
            text = line.strip()
            text_width = c.stringWidth(text, "Helvetica-Bold", 14)
            c.drawString((width - text_width) / 2, y, text)
            y -= 22
            c.setFont(font_name, 12)
            continue
        # Wrap line if too long
        words = line.split(' ')
        current_line = ''
        for word in words:
            test_line = (current_line + ' ' + word).strip()
            if c.stringWidth(test_line, font_name, 12) > max_width:
                if y < 60:
                    c.showPage()
                    c.setFont(font_name, 12)
                    y = height - 40
                c.drawString(30, y, current_line)
                y -= 16
                current_line = word
            else:
                current_line = test_line
        if current_line:
            if y < 60:
                c.showPage()
                c.setFont(font_name, 12)
                y = height - 40
            c.drawString(30, y, current_line)
            y -= 16
    c.save()
    buffer.seek(0)
    return buffer

def create_docx_from_inputs_preserve_format(raw_text, file_title="Question Paper"): 
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    doc = Document()
    # Title
    title = doc.add_paragraph()
    run = title.add_run(file_title)
    run.bold = True
    run.font.size = Pt(18)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Add a blank line
    doc.add_paragraph("")
    # Add the rest of the content, preserving formatting
    lines = raw_text.replace('**', '').replace('##', '').split('\n')
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line.replace('\t', '    '))
        run.font.size = Pt(12)
    return doc

def init_db():
    conn = sqlite3.connect('users.db')
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS users (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT UNIQUE NOT NULL,
        password TEXT NOT NULL
    )''')
    conn.commit()
    conn.close()

def register_user(username, password):
    conn = sqlite3.connect('users.db')
    c = conn.cursor()
    try:
        c.execute('INSERT INTO users (username, password) VALUES (?, ?)', (username, password))
        conn.commit()
        conn.close()
        return True, "Registration successful!"
    except sqlite3.IntegrityError:
        conn.close()
        return False, "Username already exists."

def authenticate_user(username, password):
    conn = sqlite3.connect('users.db')
    c = conn.cursor()
    c.execute('SELECT * FROM users WHERE username=? AND password=?', (username, password))
    user = c.fetchone()
    conn.close()
    return user is not None

def show_login_page():
    st.markdown('<div class="main-title">Login</div>', unsafe_allow_html=True)
    username = st.text_input("Username")
    password = st.text_input("Password", type="password")
    if st.button("Login"):
        if authenticate_user(username, password):
            st.success("Login successful!")
            st.session_state["logged_in"] = True
            st.session_state["username"] = username
            st.experimental_rerun()
        else:
            st.error("Invalid username or password.")
    st.markdown("<br>Don't have an account? <a href='#' id='show-register'>Register here</a>", unsafe_allow_html=True)
    st.markdown("""
        <script>
        const reg = document.getElementById('show-register');
        if (reg) reg.onclick = function() { window.location.hash = '#register'; window.location.reload(); };
        </script>
    """, unsafe_allow_html=True)

def show_register_page():
    st.markdown('<div class="main-title">Register</div>', unsafe_allow_html=True)
    new_username = st.text_input("Choose a Username", key="reg_user")
    new_password = st.text_input("Choose a Password", type="password", key="reg_pass")
    confirm_password = st.text_input("Confirm Password", type="password", key="reg_conf")
    if st.button("Register"):
        if not new_username or not new_password:
            st.warning("Please fill all fields.")
        elif new_password != confirm_password:
            st.error("Passwords do not match.")
        else:
            success, msg = register_user(new_username, new_password)
            if success:
                st.success(msg)
                st.session_state["show_register"] = False
                st.experimental_rerun()
            else:
                st.error(msg)
    st.markdown("<br>Already have an account? <a href='#' id='show-login'>Login here</a>", unsafe_allow_html=True)
    st.markdown("""
        <script>
        const log = document.getElementById('show-login');
        if (log) log.onclick = function() { window.location.hash = ''; window.location.reload(); };
        </script>
    """, unsafe_allow_html=True)

def main():
    init_db()
    st.markdown(
        '''
        <style>
        body {
            background: linear-gradient(120deg, #e0eafc 0%, #cfdef3 100%);
        }
        .stApp {
            background: linear-gradient(120deg, #e0eafc 0%, #cfdef3 100%);
        }
        .main-title {
            font-size: 2.5rem;
            font-weight: bold;
            color: #003366;
            text-align: center;
            margin-bottom: 0.5em;
            letter-spacing: 2px;
        }
        .subtitle {
            font-size: 1.2rem;
            color: #444;
            text-align: center;
            margin-bottom: 2em;
        }
        .block-container {
            background: #fff;
            border-radius: 18px;
            box-shadow: 0 4px 24px rgba(0,0,0,0.08);
            padding: 2.5em 2em 2em 2em;
            margin-top: 2em;
        }
        label, .stTextInput label, .stTextInput label span, .stTextInput>div>div>input, .stTextInput>div>div>div>input {
            color: #222 !important;
            font-weight: 600;
        }
        .stButton>button {
            background: linear-gradient(90deg, #003366 0%, #005fa3 100%);
            color: #fff;
            border-radius: 8px;
            font-weight: bold;
            font-size: 1.1rem;
            padding: 0.6em 2em;
            border: none;
            margin-top: 1em;
        }
        .stTextInput>div>div>input {
            border-radius: 8px;
            border: 1.5px solid #003366;
            font-size: 1.1rem;
            color: #222;
        }
        .stNumberInput>div>input {
            border-radius: 8px;
            border: 1.5px solid #003366;
            font-size: 1.1rem;
            color: #222;
        }
        .stFileUploader>div>div {
            border-radius: 8px;
            border: 1.5px solid #003366;
        }
        .logo {
            display: flex;
            justify-content: center;
            align-items: center;
            margin-bottom: 1.5em;
        }
        .logo img {
            max-width: 350px;
            height: auto;
        }
        </style>
        ''', unsafe_allow_html=True
    )
    st.markdown(
        '<div class="logo">'
        '<img src="app/static/logo.png" alt="LJ University Logo" />'
        '</div>'
        '<div class="main-title">📝 Question Paper Generator</div>'
        '<div class="subtitle">Generate beautiful, exam-ready question papers from your own study material using AI!</div>',
        unsafe_allow_html=True
    )
    st.title("Question Paper Generator")
    st.write("Upload your study material and generate a question paper based on your parameters.")
    tab1, tab2, tab3 = st.tabs(["AI Generation", "Manual Input", "Paste Formatted Paper"])
    with tab1:
        api_key = st.text_input("Enter your Gemini API Key", type="password")
        uploaded_file = st.file_uploader("Upload Material (PDF, DOCX, or TXT)", type=["pdf", "docx", "txt"])
        subject = st.text_input("Enter Subject Name", value="Fundamentals of AI")
        total_marks = st.number_input("Total Marks", min_value=1, value=50)
        mcq_count = st.number_input("Number of 1-mark MCQs", min_value=0, value=5)
        one_liner_count = st.number_input("Number of 1-mark One-liners", min_value=0, value=5)
        short_count = st.number_input("Number of 2-mark Questions", min_value=0, value=5)
        long_count = st.number_input("Number of 5-mark Questions", min_value=0, value=2)
        if st.button("Generate Question Paper", key="ai_gen_btn"):
            if not api_key:
                st.error("Please enter your Gemini API key.")
                return
            if not uploaded_file:
                st.error("Please upload your study material.")
                return
            # Extract text
            if uploaded_file.type == "application/pdf":
                material = extract_text_from_pdf(uploaded_file)
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                material = extract_text_from_docx(uploaded_file)
            else:
                material = uploaded_file.read().decode("utf-8")
            params = {
                "total_marks": total_marks,
                "mcq_count": mcq_count,
                "one_liner_count": one_liner_count,
                "short_count": short_count,
                "long_count": long_count
            }
            with st.spinner("Generating questions..."):
                questions = generate_questions_gemini(api_key, material, params)
            # Remove common AI intro lines and extra headers
            if questions:
                lines = questions.split('\n')
                cleaned_lines = []
                skip_patterns = [
                    r"^okay,? here is.*question paper.*$",
                    r"^here is a question paper based on the provided material.*$",
                    r"^lj polytechnic$",
                    r"^fundamentals of ai.*question paper$",
                    r"^subject:.*$",
                    r"^exam title:.*$",
                    r"^total marks:.*$",
                    r"^time:.*$"
                ]
                import re
                for line in lines:
                    l = line.strip().lower()
                    if not l:
                        continue  # skip blank lines
                    if any(re.match(pat, l) for pat in skip_patterns):
                        continue
                    cleaned_lines.append(line)
                questions = '\n'.join(cleaned_lines)
            st.subheader("Generated Question Paper:")
            st.text_area("Questions", questions, height=300)
            if not questions or questions.strip() == "" or questions.lower().startswith("error"):
                st.warning("No questions generated or an error occurred. Please check your API key, material, and try again.")
                return
            pdf_file = create_pdf(questions, total_marks, subject=subject)
            st.download_button(
                label="Download as PDF (Exam Style)",
                data=pdf_file,
                file_name=f"question_paper_{datetime.date.today().strftime('%Y%m%d')}.pdf",
                mime="application/pdf"
            )
            # New: Download as PDF (Preserve Format)
            pdf_preserve = create_pdf_from_inputs_preserve_format(questions, file_title=subject)
            st.download_button(
                label="Download as PDF (Preserve Format)",
                data=pdf_preserve,
                file_name=f"question_paper_{datetime.date.today().strftime('%Y%m%d')}_preserve.pdf",
                mime="application/pdf"
            )
    with tab2:
        st.markdown("**Manual Input (Paste your full formatted question paper below. All formatting is preserved.)**")
        manual_text = st.text_area("Manual Question Paper Input", height=800, key="manual_full_text")
        manual_title = st.text_input("Title (for header)", value="Question Paper", key="manual_full_title")
        if st.button("Generate PDF from Manual Input"):
            if not manual_text.strip():
                st.warning("Please paste your question paper.")
            else:
                pdf_file = create_pdf_from_inputs_preserve_format(manual_text, file_title=manual_title)
                st.download_button(
                    label="Download as PDF",
                    data=pdf_file,
                    file_name=f"question_paper_{datetime.date.today().strftime('%Y%m%d')}_manual.pdf",
                    mime="application/pdf"
                )
    with tab3:
        st.markdown("**Paste your formatted question paper below. All formatting is preserved.**")
        raw_text = st.text_area("Paste Formatted Paper", height=800, key="formatted_paper")
        file_title = st.text_input("Title (for header)", value="Question Paper", key="formatted_title")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Generate PDF from Formatted Input"):
                if not raw_text.strip():
                    st.warning("Please paste your formatted question paper.")
                else:
                    pdf_file = create_pdf_from_inputs_preserve_format(raw_text, file_title=file_title)
                    st.download_button(
                        label="Download as PDF",
                        data=pdf_file,
                        file_name=f"question_paper_{datetime.date.today().strftime('%Y%m%d')}_formatted.pdf",
                        mime="application/pdf"
                    )
        with col2:
            if st.button("Generate DOCX from Formatted Input"):
                if not raw_text.strip():
                    st.warning("Please paste your formatted question paper.")
                else:
                    docx_obj = create_docx_from_inputs_preserve_format(raw_text, file_title=file_title)
                    docx_buffer = io.BytesIO()
                    docx_obj.save(docx_buffer)
                    docx_buffer.seek(0)
                    st.download_button(
                        label="Download as DOCX",
                        data=docx_buffer,
                        file_name=f"question_paper_{datetime.date.today().strftime('%Y%m%d')}_formatted.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

if __name__ == "__main__":
    main()
